from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def generar_informe(
    metricas: dict,
    diagnostico: dict,
    config: dict,
    output_path: str = "informe_rotor.docx"
):
    doc = Document()

    # --- ESTILOS GLOBALES ---
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # --- CABECERA ---
    titulo = doc.add_heading("INFORME TÉCNICO DE ANÁLISIS DE ROTOR", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    subtitulo = doc.add_paragraph("ADF_Blades — Consultoría Técnica UAV")
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    subtitulo.runs[0].font.size = Pt(10)

    doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}")
    doc.add_paragraph("─" * 60)

    # --- CONFIGURACIÓN DEL DRON ---
    doc.add_heading("1. Configuración del sistema", level=2)

    tabla_config = doc.add_table(rows=1, cols=2)
    tabla_config.style = "Table Grid"
    hdr = tabla_config.rows[0].cells
    hdr[0].text = "Parámetro"
    hdr[1].text = "Valor"

    for key, value in config.items():
        row = tabla_config.add_row().cells
        row[0].text = str(key)
        row[1].text = str(value)

    doc.add_paragraph("")

    # --- MÉTRICAS CALCULADAS ---
    doc.add_heading("2. Métricas calculadas", level=2)

    tabla_met = doc.add_table(rows=1, cols=3)
    tabla_met.style = "Table Grid"
    hdr2 = tabla_met.rows[0].cells
    hdr2[0].text = "Métrica"
    hdr2[1].text = "Valor"
    hdr2[2].text = "Unidad"

    metricas_labels = {
        "tip_speed_m_s":         ("Tip Speed",              "m/s"),
        "mach_tip":              ("Mach Tip",               "—"),
        "disk_loading_N_m2":     ("Disk Loading",           "N/m²"),
        "CT_calculado":          ("Coef. Empuje CT",        "—"),
        "thrust_por_rotor_N":    ("Thrust por rotor",       "N"),
        "potencia_ideal_hover_W":("Potencia ideal hover",   "W"),
        "rho_kg_m3":             ("Densidad del aire",      "kg/m³"),
    }

    for key, (label, unidad) in metricas_labels.items():
        if key in metricas:
            row = tabla_met.add_row().cells
            row[0].text = label
            row[1].text = str(metricas[key])
            row[2].text = unidad

    doc.add_paragraph("")

    # --- DIAGNÓSTICO ---
    doc.add_heading("3. Diagnóstico", level=2)

    severidad = diagnostico["severidad_global"].upper()
    p_sev = doc.add_paragraph(f"Severidad global: {severidad}")
    p_sev.runs[0].bold = True
    if severidad == "CRITICAL":
        p_sev.runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    elif severidad == "WARNING":
        p_sev.runs[0].font.color.rgb = RGBColor(0xFF, 0x80, 0x00)
    else:
        p_sev.runs[0].font.color.rgb = RGBColor(0x00, 0x80, 0x00)

    doc.add_paragraph("")

    tabla_flags = doc.add_table(rows=1, cols=3)
    tabla_flags.style = "Table Grid"
    hdr3 = tabla_flags.rows[0].cells
    hdr3[0].text = "Código"
    hdr3[1].text = "Severidad"
    hdr3[2].text = "Descripción"

    for flag in diagnostico["flags"]:
        row = tabla_flags.add_row().cells
        row[0].text = flag["codigo"]
        row[1].text = flag["severidad"].upper()
        row[2].text = flag["mensaje"]

    doc.add_paragraph("")

    # --- RECOMENDACIONES ---
    doc.add_heading("4. Recomendaciones", level=2)

    for i, rec in enumerate(diagnostico["recomendaciones"], 1):
        doc.add_paragraph(f"{i}. {rec}", style="List Number")

    doc.add_paragraph("")

    # --- NOTAS DEL INGENIERO ---
    doc.add_heading("5. Notas del ingeniero", level=2)
    doc.add_paragraph("[Espacio reservado para anotaciones manuales]")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")

    # --- PIE ---
    doc.add_paragraph("─" * 60)
    pie = doc.add_paragraph("Documento generado automáticamente · Motor v0.1")
    pie.runs[0].font.size = Pt(9)
    pie.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)
    print(f"\nInforme guardado en: {output_path}")